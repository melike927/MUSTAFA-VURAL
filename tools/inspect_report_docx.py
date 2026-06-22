from docx import Document

p = r"c:\Users\melike\AppData\Local\Packages\5319275A.WhatsAppDesktop_cv1g1gvanyjgm\LocalState\sessions\A2C42D50E1A26C42297E87442331A5C168642750\transfers\2026-14\Faaliyet Raporu 2025.docx"
doc = Document(p)
print('paras', len(doc.paragraphs), 'tables', len(doc.tables))

for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f'P{idx+1}: {text}')
        if 'Modül 2' in text or 'Modul 2' in text:
            print('--- around module 2 ---')
            for j in range(max(0, idx-3), min(len(doc.paragraphs), idx+15)):
                t = doc.paragraphs[j].text.strip()
                if t:
                    print(f'P{j+1}: {t}')
            break

for i, t in enumerate(doc.tables, 1):
    print(f'TABLE {i}: rows={len(t.rows)} cols={len(t.columns)}')
    for r in range(min(8, len(t.rows))):
        print([c.text for c in t.rows[r].cells])
    print('---')
