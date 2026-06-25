from datetime import datetime
from docx import Document
from openpyxl import load_workbook

excel_path = r"c:\Users\melike\Desktop\İş Destek Verilenler.xlsx"
word_path = r"c:\Users\melike\Desktop\2026-MODUL.docx"
out_path = r"c:\Users\melike\Desktop\2026-MODUL_mart_eklendi.docx"


def s(v):
    if v is None:
        return ""
    return str(v).strip()


wb = load_workbook(excel_path, data_only=True)
ws = wb.active
headers = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
idx = {h: i + 1 for i, h in enumerate(headers)}

records = []
for r in range(2, ws.max_row + 1):
    dt = ws.cell(r, idx["Destek Tarihi"]).value
    if not isinstance(dt, datetime):
        continue
    if dt.year == 2026 and dt.month == 3:
        records.append(
            {
                "tarih": dt,
                "unvan": s(ws.cell(r, idx["Unvan"]).value),
                "ad": s(ws.cell(r, idx["Adı Soyadı"]).value),
                "fakulte": s(ws.cell(r, idx["Fakülte"]).value),
                "bolum": s(ws.cell(r, idx["Bölüm"]).value),
                "kurum": s(ws.cell(r, idx["Proje Kurumu"]).value),
                "kod": s(ws.cell(r, idx["Proje Kodu"]).value),
                "proje": s(ws.cell(r, idx["Proje Adı"]).value),
                "aciklama": s(ws.cell(r, idx["Açıklama"]).value),
            }
        )

records.sort(key=lambda x: x["tarih"])

doc = Document(word_path)

# Pick styles from existing content to keep visual consistency.
heading_style = doc.paragraphs[0].style if doc.paragraphs else None
body_style = None
for p in doc.paragraphs:
    t = p.text.strip()
    if t and "tarihinde" in t:
        body_style = p.style
        break

# Avoid duplicate insertion if already exists.
existing_text = "\n".join(p.text for p in doc.paragraphs)
if "Mart Ayı Yapılan Çalışmalar" in existing_text:
    print("Mart section already exists. No changes made.")
    doc.save(out_path)
    print("saved", out_path)
    print("count", len(records))
    raise SystemExit(0)

# Spacer
if doc.paragraphs and doc.paragraphs[-1].text.strip() != "":
    doc.add_paragraph("")

heading = "01.03.2026-31.03.2026 Mart Ayı Yapılan Çalışmalar:"
p_head = doc.add_paragraph(heading)
if heading_style is not None:
    p_head.style = heading_style

doc.add_paragraph("")

for rec in records:
    tarih = rec["tarih"].strftime("%d.%m.%Y") if hasattr(rec["tarih"], "strftime") else ""
    if tarih.startswith("0"):
        tarih = tarih[1:]

    org = ", ".join([x for x in [rec["fakulte"], rec["bolum"]] if x])
    kimlik = " ".join([x for x in [rec["unvan"], rec["ad"]] if x]).strip()
    proje_kismi = " ".join([x for x in [rec["kurum"], rec["kod"], rec["proje"]] if x]).strip()

    if org and kimlik:
        prefix = f"{tarih} tarihinde, {org} {kimlik}"
    elif kimlik:
        prefix = f"{tarih} tarihinde, {kimlik}"
    else:
        prefix = f"{tarih} tarihinde"

    sentence = prefix
    if proje_kismi:
        sentence += f" {proje_kismi}"
    if rec["aciklama"]:
        sentence += f" {rec['aciklama']}"

    p = doc.add_paragraph(sentence)
    if body_style is not None:
        p.style = body_style

doc.save(out_path)
print("saved", out_path)
print("count", len(records))
