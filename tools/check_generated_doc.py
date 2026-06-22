from docx import Document

p = r"c:\Users\melike\Desktop\MUSTAFA VURAL\Modul_2_2026_Ilk_Uc_Ay_Metin_Format.docx"
d = Document(p)
print('paras', len(d.paragraphs), 'tables', len(d.tables))
count = 0
for i, para in enumerate(d.paragraphs, 1):
    t = para.text.strip()
    if t:
        print(i, t)
        count += 1
    if count >= 15:
        break
