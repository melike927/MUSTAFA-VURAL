from docx import Document
from docx.shared import RGBColor

path = r"c:\Users\melike\Desktop\2026 modul.docx"


def normalize_run(run):
    # Force readable text styling.
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.color.theme_color = None
    run.font.hidden = False
    run.font.shadow = False


def process_paragraphs(paragraphs):
    count = 0
    for p in paragraphs:
        for run in p.runs:
            normalize_run(run)
            count += 1
    return count


doc = Document(path)
changed = 0

changed += process_paragraphs(doc.paragraphs)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            changed += process_paragraphs(cell.paragraphs)

# Also process headers/footers in all sections.
for section in doc.sections:
    changed += process_paragraphs(section.header.paragraphs)
    changed += process_paragraphs(section.footer.paragraphs)

out_path = path
backup_path = path.replace('.docx', '_yedek.docx')
fixed_path = path.replace('.docx', '_duzeltildi.docx')

# Save backup first, then overwrite original as requested.
Document(path).save(backup_path)
try:
    doc.save(out_path)
    saved_to = out_path
except PermissionError:
    doc.save(fixed_path)
    saved_to = fixed_path

print('backup', backup_path)
print('updated', saved_to)
print('runs_processed', changed)
