from docx import Document
p = r"c:\Users\melike\Desktop\2026-MODUL.docx"
d = Document(p)
print('paragraphs', len(d.paragraphs), 'tables', len(d.tables))
for i, para in enumerate(d.paragraphs, 1):
    t = para.text.strip()
    if t:
        print(f"{i}: {t}")
