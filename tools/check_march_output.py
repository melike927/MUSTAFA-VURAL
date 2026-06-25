from docx import Document
p = r"c:\Users\melike\Desktop\2026-MODUL_mart_eklendi.docx"
d = Document(p)
print('paragraphs', len(d.paragraphs), 'tables', len(d.tables))
for i, para in enumerate(d.paragraphs, 1):
    t = para.text.strip()
    if 'Mart Ayı Yapılan Çalışmalar' in t:
        print('heading_line', i, t)
        for j in range(i, min(i + 8, len(d.paragraphs))):
            tx = d.paragraphs[j].text.strip()
            if tx:
                print(j + 1, tx)
        break