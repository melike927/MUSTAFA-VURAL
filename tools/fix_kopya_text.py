from docx import Document
from docx.shared import RGBColor

path = r"c:\Users\melike\Desktop\2026-MODUL_mart_eklendi - Kopya.docx"


def normalize_run(run):
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.color.theme_color = None
    run.font.hidden = False
    run.font.shadow = False


def process_paragraphs(paragraphs):
    count = 0
    for p in paragraphs:
        for run in p.runs:
            normalize_run(run)
            count += 1
    return count


doc = Document(path)
changed = 0

changed += process_paragraphs(doc.paragraphs)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            changed += process_paragraphs(cell.paragraphs)

for section in doc.sections:
    changed += process_paragraphs(section.header.paragraphs)
    changed += process_paragraphs(section.footer.paragraphs)

out_path = path.replace('.docx', '_fixed.docx')
try:
    doc.save(path)
    out_path = path
except PermissionError:
    doc.save(out_path)

print('fixed', out_path)
print('runs_processed', changed)
