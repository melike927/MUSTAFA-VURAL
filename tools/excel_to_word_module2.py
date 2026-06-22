from datetime import datetime
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

excel_path = r"c:\Users\melike\AppData\Local\Packages\5319275A.WhatsAppDesktop_cv1g1gvanyjgm\LocalState\sessions\A2C42D50E1A26C42297E87442331A5C168642750\transfers\2026-14\İş Destek Verilenler_2026_ilk_uc_ay.xlsx"
word_path = r"c:\Users\melike\AppData\Local\Packages\5319275A.WhatsAppDesktop_cv1g1gvanyjgm\LocalState\sessions\A2C42D50E1A26C42297E87442331A5C168642750\transfers\2026-14\Modul_2_2026_Ilk_Uc_Ay_Listesi.docx"

wb = load_workbook(excel_path, data_only=True)
ws = wb.active

headers = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
rows = []
for r in range(2, ws.max_row + 1):
    row_vals = []
    empty = True
    for c in range(1, ws.max_column + 1):
        v = ws.cell(r, c).value
        if v is not None and str(v).strip() != "":
            empty = False
        if isinstance(v, datetime):
            v = v.strftime("%d.%m.%Y")
        row_vals.append("" if v is None else str(v))
    if not empty:
        rows.append(row_vals)

doc = Document()

title = doc.add_paragraph("Modül 2 - 2026 İlk Üç Ay İş Destek Verilenler Listesi")
title_format = title.runs[0].font
title_format.size = Pt(14)
title_format.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

info = doc.add_paragraph(f"Toplam kayıt: {len(rows)}")
info.runs[0].font.size = Pt(10)

# Excel-like table
table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
table.style = "Table Grid"

# Header row
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = "" if h is None else str(h)
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)

# Data rows
for r_idx, row in enumerate(rows, start=1):
    for c_idx, value in enumerate(row):
        cell = table.cell(r_idx, c_idx)
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)

# Set reasonable margins for wide table
section = doc.sections[0]
section.left_margin = Pt(20)
section.right_margin = Pt(20)
section.top_margin = Pt(25)
section.bottom_margin = Pt(25)

doc.save(word_path)
print("saved", word_path)
print("rows", len(rows))
