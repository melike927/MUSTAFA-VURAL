import calendar
from collections import defaultdict
from datetime import datetime

from docx import Document
from openpyxl import load_workbook

excel_path = r"c:\Users\melike\AppData\Local\Packages\5319275A.WhatsAppDesktop_cv1g1gvanyjgm\LocalState\sessions\A2C42D50E1A26C42297E87442331A5C168642750\transfers\2026-14\İş Destek Verilenler_2026_ilk_uc_ay.xlsx"
template_path = r"c:\Users\melike\Desktop\2026 modul.docx"
out_path = r"c:\Users\melike\Desktop\MUSTAFA VURAL\Modul_2_2026_Ilk_Uc_Ay_Metin_Format.docx"

month_names = {
    1: "Ocak",
    2: "Şubat",
    3: "Mart",
}


def s(value):
    if value is None:
        return ""
    return str(value).strip()

wb = load_workbook(excel_path, data_only=True)
ws = wb.active
headers = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
idx = {h: i + 1 for i, h in enumerate(headers)}

records = []
for r in range(2, ws.max_row + 1):
    dt = ws.cell(r, idx["Destek Tarihi"]).value
    if not isinstance(dt, datetime):
        continue
    if dt.year != 2026 or dt.month not in (1, 2, 3):
        continue
    rec = {
        "tarih": dt,
        "unvan": s(ws.cell(r, idx["Unvan"]).value),
        "ad": s(ws.cell(r, idx["Adı Soyadı"]).value),
        "fakulte": s(ws.cell(r, idx["Fakülte"]).value),
        "bolum": s(ws.cell(r, idx["Bölüm"]).value),
        "kurum": s(ws.cell(r, idx["Proje Kurumu"]).value),
        "kod": s(ws.cell(r, idx["Proje Kodu"]).value),
        "proje": s(ws.cell(r, idx["Proje Adı"]).value),
        "durum": s(ws.cell(r, idx["Durumu"]).value),
        "aciklama": s(ws.cell(r, idx["Açıklama"]).value),
    }
    records.append(rec)

records.sort(key=lambda x: x["tarih"])
by_month = defaultdict(list)
for rec in records:
    by_month[rec["tarih"].month].append(rec)

# Use template styles by reading it, then creating a new document with same defaults.
tpl = Document(template_path)
heading_style = tpl.paragraphs[0].style if tpl.paragraphs else None
text_style = None
for p in tpl.paragraphs:
    if p.text.strip() and "tarihinde" in p.text:
        text_style = p.style
        break

doc = Document()

for month in (1, 2, 3):
    last_day = calendar.monthrange(2026, month)[1]
    head = f"01.{month:02d}.2026-{last_day:02d}.{month:02d}.2026 {month_names[month]} Ayı Yapılan Çalışmalar:"
    p_head = doc.add_paragraph(head)
    if heading_style is not None:
        p_head.style = heading_style

    month_rows = by_month.get(month, [])
    if not month_rows:
        p_none = doc.add_paragraph("Bu ay için kayıt bulunmamaktadır.")
        if text_style is not None:
            p_none.style = text_style
        doc.add_paragraph("")
        continue

    for i, rec in enumerate(month_rows, start=1):
        tarih = rec["tarih"].strftime("%d.%m.%Y")

        org_parts = [x for x in [rec["fakulte"], rec["bolum"]] if x]
        org_text = ", ".join(org_parts)
        if org_text:
            org_text += ", "

        kod_kurum = " - ".join([x for x in [rec["kod"], rec["kurum"]] if x])
        if kod_kurum:
            kod_kurum += " başvurusunun "

        durum_text = rec["durum"].lower() if rec["durum"] else "destek sağlanmıştır"
        if not durum_text.endswith("."):
            durum_text += "."

        proje_adi = rec["proje"] if rec["proje"] else "ilgili proje"
        line = (
            f"{i}. {tarih} tarihinde, {org_text}{rec['unvan']} {rec['ad']}e "
            f"{kod_kurum}{proje_adi} projesi için {durum_text}"
        )

        if rec["aciklama"]:
            line += f" Açıklama: {rec['aciklama']}"

        p = doc.add_paragraph(line)
        if text_style is not None:
            p.style = text_style

        doc.add_paragraph("")

    doc.add_paragraph("")

# Remove accidental trailing empty paragraph excess isn't critical.
doc.save(out_path)
print("saved", out_path)
print("total_records", len(records))
for m in (1, 2, 3):
    print("month", m, "count", len(by_month.get(m, [])))
