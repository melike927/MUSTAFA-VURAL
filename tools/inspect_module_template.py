from docx import Document

p = r"c:\Users\melike\Desktop\2026 modul.docx"
doc = Document(p)
print('paragraphs', len(doc.paragraphs), 'tables', len(doc.tables))

for i, para in enumerate(doc.paragraphs[:120], 1):
    t = para.text.strip()
    if t:
        print(f'P{i}: {t}')

for ti, table in enumerate(doc.tables, 1):
    print(f'--- TABLE {ti} rows={len(table.rows)} cols={len(table.columns)} ---')
    max_rows = min(12, len(table.rows))
    for r in range(max_rows):
        vals = [c.text.replace('\n', ' ').strip() for c in table.rows[r].cells]
        print(vals)
